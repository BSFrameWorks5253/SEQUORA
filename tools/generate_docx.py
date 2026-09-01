import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(40, 40, 40)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # Title Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_icon = title_p.add_run("SEQUORA STUDIO\n")
    run_icon.font.name = 'Calibri'
    run_icon.font.size = Pt(11)
    run_icon.font.bold = True
    run_icon.font.color.rgb = RGBColor(120, 90, 220)

    run_title = title_p.add_run("Dedicated 1 TB+ Uploader Architecture & Guide")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(20, 25, 45)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("A complete technical guide explaining the mechanisms, architecture, and production implementation for uploading 1 TB+ of Master 4K Videos and Photo Archives to Google Drive with 100% crash-proof reliability.")
    run_sub.font.size = Pt(10.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(90, 95, 110)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    r_div = p_div.add_run("―" * 58)
    r_div.font.color.rgb = RGBColor(210, 215, 230)

    def add_section_header(title_text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(title_text)
        r.font.name = 'Calibri'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 45, 90)

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = RGBColor(30, 30, 30)
        r2 = p.add_run(text)
        r2.font.color.rgb = RGBColor(50, 50, 50)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.8)
        set_cell_background(cell, "F4F5F9")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(code_text)
        r.font.name = 'Consolas'
        r.font.size = Pt(9.0)
        r.font.color.rgb = RGBColor(35, 45, 65)
        
        # Add space after table
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(6)

    def add_callout(text, title="NOTE"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.8)
        set_cell_background(cell, "FEF3C7")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r_head = p.add_run(f"⚠️ {title}: ")
        r_head.bold = True
        r_head.font.color.rgb = RGBColor(146, 64, 14)
        r_body = p.add_run(text)
        r_body.font.color.rgb = RGBColor(120, 53, 15)
        
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(6)

    # 1. Why 1 TB Uploads Break
    add_section_header("1. Why 1 TB Uploads Break in Web Browsers & Apps Script")
    
    # Table Comparison
    tbl = doc.add_table(rows=6, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Limitation Factor", "Google Apps Script / Browser", "Dedicated Python / Rclone Engine"]
    col_widths = [Inches(1.8), Inches(2.4), Inches(2.6)]
    
    for col_idx, text in enumerate(headers):
        cell = tbl.cell(0, col_idx)
        cell.width = col_widths[col_idx]
        set_cell_background(cell, "2B3A67")
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    rows_data = [
        ("Max Single File Size", "Hard cap at ~50 MB (Base64 limits)", "Supports single files up to 5 TB each"),
        ("RAM / Memory Usage", "Loads files into browser RAM (Out of Memory crash)", "Streams from disk in 16MB–64MB chunks (< 200MB RAM)"),
        ("Connection Drops", "Entire upload fails and restarts from zero", "Resumable Chunks: Resumes at exact byte where it paused"),
        ("Power Outage / Reboot", "Session lost completely", "State database remembers uploaded files and continues"),
        ("Upload Speed", "Single-thread JavaScript bottleneck", "4x to 8x Faster multi-threaded parallel transfers")
    ]

    for row_idx, data in enumerate(rows_data, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(data):
            cell = tbl.cell(row_idx, col_idx)
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(9.0)
            if col_idx == 0:
                r.font.bold = True

    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_after = Pt(8)

    # 2. Dedicated Python Uploader
    add_section_header("2. How the Dedicated Python Uploader Works")
    p_py = doc.add_paragraph()
    p_py.add_run("The Python Uploader runs natively on Windows as part of SEQUORA Studio, utilizing Google's official Drive API v3 Resumable Upload protocol.")
    
    add_code_block(
"""+-----------------------------------------------------------------------------------+
|                           SEQUORA Python Uploader                                |
+-----------------------------------------------------------------------------------+
  |
  +--> 1. Folder Scanner: Scans 1 TB local folder (5,000 photos + 20 master videos)
  |
  +--> 2. State DB (.upload_state.json): Tracks [Pending / Uploading / Completed / Failed]
  |
  +--> 3. Folder Tree Sync: Recreates directory hierarchy on Google Drive via API v3
  |
  +--> 4. Multi-Threaded Worker Pool (4 to 8 parallel streams):
  |       +-- Worker #1: 4K_Master_CamA.mp4 (64 MB chunked buffer stream via HTTP PUT)
  |       +-- Worker #2: 4K_Master_CamB.mp4 (Resumable session URI from Google API)
  |       +-- Worker #3: RAW_Photos_Batch1.zip (Direct stream)
  |       +-- Worker #4: RAW_Photos_Batch2.zip (Direct stream)
  |
  +--> 5. Auto-Retry & Backoff: Catches HTTP 429/503 and auto-retries seamlessly
  |
  +--> 6. SHA-256 Checksum: Verifies zero corrupted frames or missing bytes"""
    )

    add_bullet("Google Drive Resumable Protocol: ", "Files are read from the hard drive in 64 MB buffer chunks. Python requests a unique session URI and streams chunks via HTTP PUT. If interrupted at 92%, it queries Google Drive for received bytes and resumes from the 93rd percent.")
    add_bullet("Persistent State Database: ", "Progress is recorded to a local state file (.upload_state.json). If your computer loses power or reboots, launching the script automatically skips finished files and resumes the unfinished ones.")
    add_bullet("MD5 / SHA-256 Verification: ", "Google Drive returns an MD5 hash upon completion. Python validates this hash against the local file to ensure 100% data integrity.")

    # 3. Dedicated Rclone Uploader
    add_section_header("3. How the Dedicated Rclone Uploader Works")
    p_rc = doc.add_paragraph()
    p_rc.add_run("Rclone is an industry-standard, high-performance cloud sync engine written in Go. It provides maximum bandwidth utilization and resilience for large studio archives.")

    add_code_block(
"""Command Example:
rclone copy "D:\\Event_Master_1TB" "MainDrive:SEQUORA_Master_Archives/Event_Master" ^
  --transfers 4 ^
  --checkers 8 ^
  --drive-chunk-size 64M ^
  --fast-list ^
  --progress ^
  --retries 5"""
    )

    add_bullet("Raw Socket Performance: ", "Written in Go with direct kernel socket streaming, saturating 100% of your fiber internet bandwidth.")
    add_bullet("Drive Chunk Size Tuning: ", "--drive-chunk-size 64M or 128M optimizes throughput for 10 GB - 100 GB 4K video files.")
    add_bullet("Bandwidth Scheduling: ", "Allows bandwidth throttling during studio office hours and full speed overnight.")

    # 4. Comparison Table
    add_section_header("4. Method Comparison: Python vs Rclone vs Drive Desktop")
    
    tbl2 = doc.add_table(rows=7, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    heads2 = ["Criteria", "Google Drive Desktop", "Rclone Engine", "Custom Python Engine"]
    col_w2 = [Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.7)]

    for c_idx, t in enumerate(heads2):
        cell = tbl2.cell(0, c_idx)
        cell.width = col_w2[c_idx]
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=140, bottom=140, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(t)
        r.font.bold = True
        r.font.size = Pt(9.0)
        r.font.color.rgb = RGBColor(255, 255, 255)

    comp_data = [
        ("Setup Required", "Zero Setup (Windows installer)", "Single .exe portable binary", "Integrated in SEQUORA"),
        ("User Interface", "Windows File Explorer (G:\\)", "CLI / 1-Click .bat file", "Native Studio GUI"),
        ("Max File Size", "5 TB per file", "5 TB per file", "5 TB per file"),
        ("Multi-Threading", "Automatic by Google", "Configurable (4-16 streams)", "Configurable (4-8 streams)"),
        ("Auto-Resume", "Yes (Automatic)", "Yes (Automatic)", "Yes (State DB)"),
        ("Best Used For", "Zero-Code Instant Sync", "Maximum Transfer Speed", "Seamless In-App Button")
    ]

    for r_idx, row in enumerate(comp_data, start=1):
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tbl2.cell(r_idx, c_idx)
            cell.width = col_w2[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=90, bottom=90, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            if c_idx == 0:
                r.font.bold = True

    p_sp2 = doc.add_paragraph()
    p_sp2.paragraph_format.space_after = Pt(8)

    # 5. 750 GB Quota
    add_section_header("5. Google's Daily 750 GB Upload Ceiling")
    add_callout(
        "Google enforces a strict limit of 750 GB upload per 24 hours per Google account. When uploading 1 TB, the first 750 GB transfers at maximum speed. Upon hitting the quota, the engine catches error 403, pauses gracefully without losing progress, and automatically resumes the remaining 250 GB once Google resets the quota.",
        title="IMPORTANT QUOTA RULE"
    )

    # 6. Production Python Code
    add_section_header("6. Production Python Resumable Code Blueprint")
    add_code_block(
"""import os
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

def upload_large_video_resumable(drive_service, file_path, folder_id, chunk_size=64*1024*1024):
    \"\"\"
    Uploads 1GB - 50GB 4K videos in 64MB chunks without memory overload.
    \"\"\"
    file_name = os.path.basename(file_path)
    file_metadata = {'name': file_name, 'parents': [folder_id]}
    
    media = MediaFileUpload(file_path, chunksize=chunk_size, resumable=True)
    request = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id, name, size, md5Checksum'
    )
    
    response = None
    while response is None:
        status, response = request.next_chunk()
        if status:
            progress_pct = int(status.progress() * 100)
            print(f"[{file_name}] Progress: {progress_pct}%")
            
    print(f"✅ Successfully Uploaded: {file_name} (ID: {response.get('id')})")
    return response"""
    )

    # 7. Final Recommendation
    add_section_header("7. Summary & Recommended Action Plan")
    add_bullet("For 1 TB Master Videos (Account 1): ", "Use Google Drive for Desktop for instant zero-code drag-and-drop, or the Rclone 1-click batch script for maximum network saturation.")
    add_bullet("For Reference & Remaining Photos (Account 2): ", "Continue using the Turbo Google Apps Script Web Uploader (optimized for fast photo batches).")

    # Save
    out_path = os.path.abspath("docs/DEDICATED_1TB_UPLOADER_ARCHITECTURE.docx")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"Docx generated at: {out_path}")

if __name__ == "__main__":
    create_document()
